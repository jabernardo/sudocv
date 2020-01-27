# -*- coding: utf-8 -*-

from docx import Document

class Doc:
    __document = None
   
    __collector = None
    
    def __init__(self, collector):
        self.__collector = collector
        self.__document = Document()
        self.__construct()
    
    def __construct(self):
        self.__build_info()
        self.__build_projects()
        self.__document.add_page_break()
    
    def __build_info(self):
        profile = self.__collector.get_user()
        
        self.__document.add_heading(profile.name, 0)
        self.__document.add_paragraph(f'/blog:\t\t{profile.blog}')
        self.__document.add_paragraph(f'@mail:\t\t{profile.email}')
        self.__document.add_paragraph(profile.bio, style='Intense Quote')
    
    def __build_projects(self):
        self.__document.add_heading('Projects', 1)
        
        self.__document.add_heading('Repositories', 2)
        
        for proj in self.__collector.get_repositories():
            if proj.description:
                self.__document.add_heading(proj.name, 3)
                self.__document.add_paragraph(proj.description)
        
        self.__document.add_heading('Programming Languages', 2)
        
        for lang in self.__collector.get_languages():
            self.__document.add_paragraph(lang, style='List Bullet')
        
    
#    def __add_avatar(self):
#        self.__document.add_picture(self.__collector.getUser().avatar_url, width=Inches(1.25))
    
    def save(self, output = 'sample.docx'):
        self.__document.save(output)

